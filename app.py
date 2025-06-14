import os
import logging
import json
from datetime import datetime, timezone
from flask import Flask, render_template, request, redirect, url_for, flash, jsonify, make_response, session
from flask_sqlalchemy import SQLAlchemy
from sqlalchemy.orm import DeclarativeBase
from sqlalchemy import func
from werkzeug.middleware.proxy_fix import ProxyFix
import pandas as pd
import io
import re
from functools import wraps
from urllib.parse import urlparse
import uuid
import qrcode
from PIL import Image
# Import barcode libraries with correct python-barcode structure
from barcode.codex import Code128
from barcode.writer import ImageWriter
# Word document generation
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
import base64
import csv
import logging
import traceback
from datetime import datetime

# Configure logging
logging.basicConfig(level=logging.DEBUG)

# Configure barcode-specific logger
os.makedirs('logs', exist_ok=True)
barcode_logger = logging.getLogger('barcode_generation')
barcode_logger.setLevel(logging.INFO)

handler = logging.FileHandler('logs/barcode_generation.log')
formatter = logging.Formatter('%(asctime)s - %(levelname)s - %(message)s')
handler.setFormatter(formatter)
barcode_logger.addHandler(handler)

def log_barcode_operation(operation, barcode_data, barcode_type, success, error=None):
    """Log barcode generation operations with detailed information"""
    timestamp = datetime.now().isoformat()
    if success:
        barcode_logger.info(f"{operation} - {barcode_type} - {barcode_data} - SUCCESS")
    else:
        barcode_logger.error(f"{operation} - {barcode_type} - {barcode_data} - FAILED: {error}")
        barcode_logger.error(f"Stack trace: {traceback.format_exc()}")

def set_cell_border(cell):
    """Add solid black borders around a Word table cell"""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    for edge in ('top', 'left', 'bottom', 'right'):
        elm = OxmlElement(f'w:{edge}')
        elm.set(qn('w:val'), 'single')
        elm.set(qn('w:sz'), '12')  # 12 eighth-points = 1.5pt
        elm.set(qn('w:color'), '000000')  # black
        tcPr.append(elm)

def set_individual_cell_border(cell):
    """Add individual borders for each cell without sharing borders"""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    
    # Add margins for spacing between cells
    mar = OxmlElement('w:tcMar')
    for side in ('top', 'left', 'bottom', 'right'):
        side_mar = OxmlElement(f'w:{side}')
        side_mar.set(qn('w:w'), '100')  # 100 twips spacing
        side_mar.set(qn('w:type'), 'dxa')
        mar.append(side_mar)
    tcPr.append(mar)
    
    # Add thick borders around each cell
    for edge in ('top', 'left', 'bottom', 'right'):
        elm = OxmlElement(f'w:{edge}')
        elm.set(qn('w:val'), 'single')
        elm.set(qn('w:sz'), '18')  # Thicker border - 18 eighth-points = 2.25pt
        elm.set(qn('w:color'), '000000')  # black
        elm.set(qn('w:space'), '0')
        tcPr.append(elm)

class Base(DeclarativeBase):
    pass

db = SQLAlchemy(model_class=Base)

# Create the app
app = Flask(__name__)
app.secret_key = os.environ.get("SESSION_SECRET", "dev-secret-key-change-in-production")
app.wsgi_app = ProxyFix(app.wsgi_app, x_proto=1, x_host=1)

# Configure the database
app.config["SQLALCHEMY_DATABASE_URI"] = os.environ.get("DATABASE_URL", "sqlite:///inventory.db")
app.config["SQLALCHEMY_ENGINE_OPTIONS"] = {
    "pool_recycle": 300,
    "pool_pre_ping": True,
}

# Initialize the app with the extension
db.init_app(app)
from sqlalchemy import text
from models import HardwareType, LotNumber, Box, PullEvent, ActionLog

# Auto-migrate on startup
with app.app_context():

    # Create tables that don’t exist yet
    db.create_all()

    # Patch the boxes table in-place
    with db.engine.begin() as conn:
        # Patch pull_events
        # Database-agnostic column checking
        try:
            dialect_name = conn.dialect.name.lower()
            if dialect_name == 'postgresql':
                result = conn.execute(text("""
                    SELECT column_name 
                    FROM information_schema.columns 
                    WHERE table_name = 'pull_events' AND table_schema = 'public'
                """))
                cols = [row[0] for row in result.fetchall()]
            elif dialect_name == 'sqlite':
                result = conn.execute(text("PRAGMA table_info(pull_events)"))
                cols = [row[1] for row in result.fetchall()]
            else:
                cols = []
        except Exception:
            cols = []

        # Rename legacy column if present
        if "quantity_pulled" in cols and "quantity" not in cols:
            conn.execute(text("ALTER TABLE pull_events RENAME COLUMN quantity_pulled TO quantity"))
            cols.append("quantity")

        if "qc_operator" in cols and "qc_personnel" not in cols:
            conn.execute(text("ALTER TABLE pull_events RENAME COLUMN qc_operator TO qc_personnel"))
            cols.append("qc_personnel")

        # Add any brand-new columns that still don’t exist
        migrations = {
            "quantity":    "ALTER TABLE pull_events ADD COLUMN quantity INTEGER DEFAULT 0",
            "mo":          "ALTER TABLE pull_events ADD COLUMN mo VARCHAR(50)",
            "operator":    "ALTER TABLE pull_events ADD COLUMN operator VARCHAR(50)",
            "qc_personnel": "ALTER TABLE pull_events ADD COLUMN qc_personnel VARCHAR(50)",
        }
        for col, ddl in migrations.items():
            if col not in cols:
                try:
                    conn.execute(text(ddl))
                except Exception as e:
                    # Column might already exist, continue
                    print(f"Migration warning for {col}: {str(e)}")
                    pass

        # Also migrate boxes table
        try:
            if dialect_name == 'postgresql':
                result = conn.execute(text("""
                    SELECT column_name 
                    FROM information_schema.columns 
                    WHERE table_name = 'boxes' AND table_schema = 'public'
                """))
                box_cols = [row[0] for row in result.fetchall()]
            elif dialect_name == 'sqlite':
                result = conn.execute(text("PRAGMA table_info(boxes)"))
                box_cols = [row[1] for row in result.fetchall()]
            else:
                box_cols = []
                
            if box_cols:
                box_migrations = {
                    "operator": "ALTER TABLE boxes ADD COLUMN operator VARCHAR(50)",
                    "qc_personnel": "ALTER TABLE boxes ADD COLUMN qc_personnel VARCHAR(50)",
                }
                for col, ddl in box_migrations.items():
                    if col not in box_cols:
                        try:
                            conn.execute(text(ddl))
                        except Exception as e:
                            print(f"Migration warning for boxes.{col}: {str(e)}")
                            pass
        except Exception as e:
            print(f"Boxes migration warning: {str(e)}")
            pass

        # Also migrate action_logs table
        try:
            if dialect_name == 'postgresql':
                result = conn.execute(text("""
                    SELECT column_name 
                    FROM information_schema.columns 
                    WHERE table_name = 'action_logs' AND table_schema = 'public'
                """))
                action_cols = [row[0] for row in result.fetchall()]
            elif dialect_name == 'sqlite':
                result = conn.execute(text("PRAGMA table_info(action_logs)"))
                action_cols = [row[1] for row in result.fetchall()]
            else:
                action_cols = []
                
            if action_cols:
                action_migrations = {
                    "qc_personnel": "ALTER TABLE action_logs ADD COLUMN qc_personnel VARCHAR(100)",
                }
                for col, ddl in action_migrations.items():
                    if col not in action_cols:
                        try:
                            conn.execute(text(ddl))
                        except Exception as e:
                            print(f"Migration warning for action_logs.{col}: {str(e)}")
                            pass
        except Exception as e:
            print(f"Action logs migration warning: {str(e)}")
            pass
            
    print("✅ Database schema auto-migration complete")


# Add JSON filter for templates
@app.template_filter('from_json')
def from_json_filter(value):
    try:
        return json.loads(value) if value else {}
    except (json.JSONDecodeError, TypeError):
        return {}

def is_safe_url(target):
    """Check if the target URL is safe for redirect"""
    if not target:
        return False
    ref_url = urlparse(request.host_url)
    test_url = urlparse(target)
    return test_url.scheme in ('http', 'https') and ref_url.netloc == test_url.netloc

def admin_required(f):
    """Decorator to require admin authentication"""
    @wraps(f)
    def decorated_function(*args, **kwargs):
        if not session.get('is_admin'):
            flash('Admin access required to manage boxes', 'error')
            next_url = request.url if is_safe_url(request.url) else url_for('manage_boxes')
            return redirect(url_for('admin_login', next=next_url))
        return f(*args, **kwargs)
    return decorated_function

def sanitize_box_id_component(component):
    """Sanitize a component for use in box ID generation"""
    # Remove special characters and replace spaces with underscores
    sanitized = re.sub(r'[^\w\-_]', '_', str(component).strip())
    # Remove multiple consecutive underscores
    sanitized = re.sub(r'_+', '_', sanitized)
    # Remove leading/trailing underscores
    return sanitized.strip('_')

def generate_box_id(hardware_type_name, lot_number_name, box_number):
    """Generate a unique box ID from components"""
    type_clean = sanitize_box_id_component(hardware_type_name)
    lot_clean = sanitize_box_id_component(lot_number_name)
    box_clean = sanitize_box_id_component(box_number)
    return f"{type_clean}_{lot_clean}_{box_clean}"

def generate_unique_barcode(length=10):
    """Generate a unique barcode checking for duplicates"""
    import random
    import string
    
    charset = string.ascii_uppercase + string.digits
    attempts = 0
    max_attempts = 1000
    
    while attempts < max_attempts:
        new_code = ''.join(random.choices(charset, k=length))
        if not Box.query.filter_by(barcode=new_code).first():
            return new_code
        attempts += 1
    
    raise ValueError("Unable to generate a unique barcode after many attempts.")

def generate_barcode_image(barcode_data, barcode_type='qrcode', format='png'):
    """Generate barcode image and save to static/barcodes/"""
    try:
        # Ensure barcodes directory exists
        barcodes_dir = os.path.join('static', 'barcodes')
        os.makedirs(barcodes_dir, exist_ok=True)
        
        # Generate filename with type suffix
        filename = f"{barcode_data}_{barcode_type}.png"
        filepath = os.path.join(barcodes_dir, filename)
        
        if barcode_type == 'code128':
            # Generate Code 128 barcode using correct python-barcode structure
            try:
                code = Code128(barcode_data, writer=ImageWriter())
                saved_filename = code.save(filepath.replace('.png', ''))  # barcode library adds .png automatically
                log_barcode_operation("GENERATE", barcode_data, barcode_type, True)
                
                # Verify file was created
                if os.path.exists(saved_filename):
                    return f"/static/barcodes/{filename}"
                else:
                    raise FileNotFoundError(f"Code128 file not created at {saved_filename}")
                    
            except Exception as e:
                error_msg = f"Code128 generation failed: {str(e)}"
                log_barcode_operation("GENERATE", barcode_data, barcode_type, False, error_msg)
                logging.error(f"Code128 generation error: {error_msg}")
                logging.error(f"Stack trace: {traceback.format_exc()}")
                # Fallback to QR code if Code 128 fails
                return generate_barcode_image(barcode_data, 'qrcode', format)
        else:
            # Generate QR code
            try:
                qr = qrcode.QRCode(
                    version=1,
                    error_correction=qrcode.constants.ERROR_CORRECT_L,
                    box_size=10,
                    border=4,
                )
                qr.add_data(barcode_data)
                qr.make(fit=True)
                
                # Create image
                img = qr.make_image(fill_color="black", back_color="white")
                img.save(filepath)
                
                # Verify file was created
                if os.path.exists(filepath):
                    log_barcode_operation("GENERATE", barcode_data, barcode_type, True)
                    return f"/static/barcodes/{filename}"
                else:
                    raise FileNotFoundError(f"QR code file not created at {filepath}")
                    
            except Exception as e:
                error_msg = f"QR code generation failed: {str(e)}"
                log_barcode_operation("GENERATE", barcode_data, barcode_type, False, error_msg)
                logging.error(f"QR code generation error: {error_msg}")
                logging.error(f"Stack trace: {traceback.format_exc()}")
                return None
        
    except Exception as e:
        error_msg = f"General barcode generation error: {str(e)}"
        log_barcode_operation("GENERATE", barcode_data, barcode_type, False, error_msg)
        logging.error(f"Barcode generation error: {error_msg}")
        logging.error(f"Stack trace: {traceback.format_exc()}")
        return None

def log_action(action_type, user, box_id=None, hardware_type=None, lot_number=None, 
               previous_quantity=None, quantity_change=None, available_quantity=None,
               operator=None, qc_personnel=None, details=None):
    """Log an admin action with enhanced tracking"""
    try:
        action_log = ActionLog()
        action_log.action_type = action_type
        action_log.user = user or operator  # Use operator as user if user not provided
        action_log.box_id = box_id
        action_log.hardware_type = hardware_type
        action_log.lot_number = lot_number
        action_log.previous_quantity = previous_quantity
        action_log.quantity_change = quantity_change
        action_log.available_quantity = available_quantity
        action_log.operator = operator
        action_log.qc_personnel = qc_personnel
        action_log.details = json.dumps(details) if details else None
        db.session.add(action_log)
        db.session.commit()
    except Exception as e:
        app.logger.error(f"Failed to log action: {str(e)}")
        # Don't fail the main operation if logging fails
        pass

def group_boxes_by_type_lot(results):
    """Group boxes by type code (first 3 digits) and type-lot combination"""
    from collections import defaultdict
    
    type_code_groups = defaultdict(lambda: defaultdict(list))
    
    for box, type_name, lot_name in results:
        # Extract first 3 digits from type name
        type_code = type_name[:3] if len(type_name) >= 3 else type_name
        # Use tuple key to avoid collisions with | character
        type_lot_key = (type_name, lot_name)
        
        type_code_groups[type_code][type_lot_key].append({
            'box': box,
            'type_name': type_name,
            'lot_name': lot_name
        })
    
    # Calculate totals for each type-lot group - convert to plain dict
    grouped_data = {}
    for type_code, type_lot_groups in type_code_groups.items():
        grouped_data[type_code] = {}
        for type_lot_key, boxes_data in type_lot_groups.items():
            type_name, lot_name = type_lot_key  # Unpack tuple
            
            total_initial = sum(bd['box'].initial_quantity for bd in boxes_data)
            total_remaining = sum(bd['box'].remaining_quantity for bd in boxes_data)
            
            # Create string key for template compatibility
            display_key = f"{type_name}__{lot_name}"  # Use __ as separator
            grouped_data[type_code][display_key] = {
                'type_name': type_name,
                'lot_name': lot_name,
                'boxes': boxes_data,
                'total_initial': total_initial,
                'total_remaining': total_remaining,
                'box_count': len(boxes_data)
            }
    
    return grouped_data

def calculate_inventory_stats(grouped_data):
    """Calculate summary statistics from grouped data with error handling"""
    total_boxes = sum(g.get('box_count', 0)
                      for lots in grouped_data.values()
                      for g in lots.values())
    
    available_boxes = 0
    empty_boxes = 0
    negative_boxes = 0  # Track anomalous negative quantities
    
    for lots in grouped_data.values():
        for g in lots.values():
            for b in g.get('boxes', []):
                remaining = b['box'].remaining_quantity
                # Handle None or unexpected values
                if remaining is None:
                    remaining = 0
                
                if remaining > 0:
                    available_boxes += 1
                elif remaining == 0:
                    empty_boxes += 1
                else:  # negative
                    empty_boxes += 1
                    negative_boxes += 1
    
    total_remaining = sum(g.get('total_remaining', 0)
                          for lots in grouped_data.values()
                          for g in lots.values())
    
    return {
        'total_boxes': total_boxes,
        'available_boxes': available_boxes,
        'empty_boxes': empty_boxes,
        'negative_boxes': negative_boxes,
        'total_remaining': max(0, total_remaining)  # Ensure non-negative
    }

def build_filtered_query(type_filter=None, lot_filter=None, search_query=None):
    """Build base query with common filters - DRY helper"""
    query = db.session.query(
        Box,
        HardwareType.name.label('type_name'),
        LotNumber.name.label('lot_name')
    ).join(HardwareType, Box.hardware_type_id == HardwareType.id)\
     .join(LotNumber, Box.lot_number_id == LotNumber.id)
    
    # Apply filters
    if type_filter:
        query = query.filter(HardwareType.name == type_filter)
    if lot_filter:
        query = query.filter(LotNumber.name == lot_filter)
    if search_query and search_query.strip():
        search_pattern = f'%{search_query.strip()}%'
        query = query.filter(
            db.or_(
                Box.box_id.ilike(search_pattern),
                Box.barcode.ilike(search_pattern),
                HardwareType.name.ilike(search_pattern),
                LotNumber.name.ilike(search_pattern)
            )
        )
    
    return query

@app.route('/')
def index():
    """Home page with navigation options"""
    return render_template('index.html')

@app.route('/admin_login', methods=['GET', 'POST'])
def admin_login():
    """Simple admin login with security"""
    if request.method == 'POST':
        username = request.form.get('username', '').strip()
        password = request.form.get('password', '')
        
        # Simple admin credentials (in production, use proper hashed passwords)
        if username == 'admin' and password == 'admin123':
            session['is_admin'] = True
            session['admin_username'] = username
            
            # Secure redirect handling
            next_url = request.args.get('next')
            if next_url and is_safe_url(next_url):
                flash('Admin login successful', 'success')
                return redirect(next_url)
            else:
                flash('Admin login successful', 'success')
                return redirect(url_for('manage_boxes'))
        else:
            flash('Invalid admin credentials', 'error')
    
    return render_template('admin_login.html')

@app.route('/admin_logout')
def admin_logout():
    """Logout admin"""
    session.pop('is_admin', None)
    session.pop('admin_username', None)
    flash('Admin logged out successfully', 'success')
    return redirect(url_for('index'))

@app.route('/add_box', methods=['GET', 'POST'])
def add_box():
    """Add a new box to inventory"""
    if request.method == 'POST':
        try:
            # Get form data
            hardware_type_name = request.form.get('hardware_type', '').strip()
            new_hardware_type = request.form.get('new_hardware_type', '').strip()
            lot_number_name = request.form.get('lot_number', '').strip()
            new_lot_number = request.form.get('new_lot_number', '').strip()
            box_number = request.form.get('box_number', '').strip()
            initial_quantity = request.form.get('initial_quantity', 0)
            barcode = request.form.get('barcode', '').strip()
            operator = request.form.get('operator', '').strip()
            qc_operator = request.form.get('qc_operator', '').strip()
            
            # Validation
            errors = []
            
            # Handle hardware type
            if new_hardware_type:
                # Check if new type already exists
                existing_type = HardwareType.query.filter_by(name=new_hardware_type).first()
                if existing_type:
                    hardware_type = existing_type
                else:
                    hardware_type = HardwareType()
                    hardware_type.name = new_hardware_type
                    db.session.add(hardware_type)
                    db.session.flush()  # Get the ID without committing
            elif hardware_type_name:
                hardware_type = HardwareType.query.filter_by(name=hardware_type_name).first()
                if not hardware_type:
                    errors.append("Invalid hardware type selected")
            else:
                errors.append("Hardware type is required")
            
            # Handle lot number
            if new_lot_number:
                # Check if new lot already exists
                existing_lot = LotNumber.query.filter_by(name=new_lot_number).first()
                if existing_lot:
                    lot_number = existing_lot
                else:
                    lot_number = LotNumber()
                    lot_number.name = new_lot_number
                    db.session.add(lot_number)
                    db.session.flush()  # Get the ID without committing
            elif lot_number_name:
                lot_number = LotNumber.query.filter_by(name=lot_number_name).first()
                if not lot_number:
                    errors.append("Invalid lot number selected")
            else:
                errors.append("Lot number is required")
            
            # Validate other fields
            if not box_number:
                errors.append("Box number is required")
            
            try:
                initial_quantity = int(initial_quantity)
                if initial_quantity <= 0:
                    errors.append("Initial quantity must be greater than 0")
            except (ValueError, TypeError):
                errors.append("Initial quantity must be a valid number")
            
            if not barcode:
                errors.append("Barcode is required")
            elif Box.query.filter_by(barcode=barcode).first():
                errors.append("Barcode already exists")
            
            if not operator:
                errors.append("Operator name is required")
            
            if not qc_operator:
                errors.append("QC Personnel name is required")
            
            if operator == qc_operator:
                errors.append("Operator and QC Personnel cannot be the same")
            
            if errors:
                for error in errors:
                    flash(error, 'error')
                # Preserve form data
                form_data = request.form.to_dict()
                types = HardwareType.query.all()
                lots = LotNumber.query.all()
                return render_template('add_box.html', types=types, lots=lots, form_data=form_data)
            
            # Generate box ID
            box_id = generate_box_id(hardware_type.name, lot_number.name, box_number)
            
            # Check if box ID already exists
            if Box.query.filter_by(box_id=box_id).first():
                flash("A box with this Type/Lot/Box combination already exists", 'error')
                form_data = request.form.to_dict()
                types = HardwareType.query.all()
                lots = LotNumber.query.all()
                return render_template('add_box.html', types=types, lots=lots, form_data=form_data)
            
            # Create new box
            new_box = Box()
            new_box.box_id = box_id
            new_box.hardware_type_id = hardware_type.id
            new_box.lot_number_id = lot_number.id
            new_box.box_number = box_number
            new_box.initial_quantity = initial_quantity
            new_box.remaining_quantity = initial_quantity
            new_box.barcode = barcode
            new_box.operator = operator
            new_box.qc_personnel = qc_operator
            
            db.session.add(new_box)
            db.session.commit()
            
            # Log the box addition
            log_action(
                action_type='box_add',
                user=operator or 'System',
                box_id=box_id,
                hardware_type=hardware_type.name,
                lot_number=lot_number.name,
                previous_quantity=0,
                quantity_change=initial_quantity,
                available_quantity=initial_quantity,
                operator=operator,
                qc_personnel=qc_operator,
                details={
                    'barcode': barcode
                }
            )
            
            flash(f"Box {box_id} added successfully!", 'success')
            return redirect(url_for('add_box'))
            
        except Exception as e:
            db.session.rollback()
            app.logger.error(f"Error adding box: {str(e)}")
            flash("An error occurred while adding the box", 'error')
    
    # GET request - show form
    types = HardwareType.query.all()
    lots = LotNumber.query.all()
    return render_template('add_box.html', types=types, lots=lots)


@app.route('/log_event', methods=['GET', 'POST'])
def log_event():
    """Log pull/return event with improved validation"""
    if request.method == 'POST':
        try:
            # Get form data
            barcode       = request.form["barcode"].strip()
            qty_str       = request.form["quantity"]
            event_type    = request.form["event_type"].lower()
            mo            = request.form["mo"].strip()
            operator      = request.form["operator"].strip()
            qc_personnel  = request.form["qc_personnel"].strip()
            signature     = request.form.get("signature", "").strip()
            
            # Validation
            errors = []
            
            if not barcode:
                errors.append("Barcode is required")
            
            try:
                quantity = int(qty_str)
                if quantity <= 0:
                    errors.append("Quantity must be greater than 0")
            except (ValueError, TypeError):
                errors.append("Quantity must be a valid number")
            
            if event_type not in ('pull', 'return'):
                errors.append("Invalid event type")
            
            if not mo:
                errors.append("Manufacturing Order (MO) is required")
            
            if not operator:
                errors.append("Operator name is required")
            
            if not qc_personnel:
                errors.append("QC Personnel name is required")
            
            if operator == qc_personnel:
                errors.append("Operator and QC Personnel cannot be the same")
            
            # Find the box
            box = None
            if barcode:
                box = Box.query.filter_by(barcode=barcode).first()
                if not box:
                    errors.append("Box with given barcode not found")
            
            if errors:
                for error in errors:
                    flash(error, 'danger')
                return redirect(url_for('log_event'))
            
            # Calculate quantity change (positive for return, negative for pull)
            change = quantity if event_type == "return" else -quantity
            new_qty = box.remaining_quantity + change

            if new_qty < 0:
                flash("Not enough quantity in box", "danger")
                return redirect(url_for('log_event'))

            # Store previous quantity for action log
            previous_qty = box.remaining_quantity
            
            # Update box quantity
            box.remaining_quantity = new_qty
            
            # Create pull event record
            pull_event = PullEvent()
            pull_event.box_id = box.id
            pull_event.quantity = change
            pull_event.mo = mo
            pull_event.operator = operator
            pull_event.qc_personnel = qc_personnel
            pull_event.signature = signature
            db.session.add(pull_event)
            
            # Create action log record with proper quantity tracking
            log_action(
                action_type        = event_type,
                user               = operator,
                box_id             = box.box_id,
                hardware_type      = box.hardware_type.name,
                lot_number         = box.lot_number.name,
                previous_quantity  = previous_qty,
                quantity_change    = change,
                available_quantity = box.remaining_quantity,
                operator           = operator,
                qc_personnel       = qc_personnel,
                details            = {"mo": mo, "signature": signature}
            )
            
            db.session.commit()
            
            flash("Event logged successfully!", "success")
            return redirect(url_for('log_event'))
            
        except Exception as e:
            db.session.rollback()
            app.logger.error(f"Error logging event: {str(e)}")
            flash("An error occurred while logging the event", 'error')
    
    return render_template('log_event.html')

@app.route('/dashboard')
def dashboard():
    """Inventory dashboard with grouped display"""
    # Get filter parameters
    type_filter = request.args.get('type_filter', '')
    lot_filter = request.args.get('lot_filter', '')
    
    # Use DRY filter helper
    query = build_filtered_query(type_filter, lot_filter)
    
    # Order by type name first, then lot name, then box number
    results = query.order_by(HardwareType.name, LotNumber.name, Box.box_number).all()
    
    # Group boxes using helper function
    grouped_data = group_boxes_by_type_lot(results)
    
    # Calculate statistics
    total_stats = calculate_inventory_stats(grouped_data)
    
    # Get unique types and lots for filter dropdowns
    types = HardwareType.query.order_by(HardwareType.name).all()
    lots = LotNumber.query.order_by(LotNumber.name).all()
    
    return render_template('dashboard.html', 
                         grouped_data=grouped_data,
                         total_stats=total_stats,
                         types=types, 
                         lots=lots,
                         type_filter=type_filter,
                         lot_filter=lot_filter)

@app.route('/box_logs/<int:box_id>')
def box_logs(box_id):
    """View logs for a specific box"""
    box = Box.query.get_or_404(box_id)
    hardware_type = HardwareType.query.get(box.hardware_type_id)
    lot_number = LotNumber.query.get(box.lot_number_id)
    
    # Get pull events for this box (box_id in PullEvent refers to Box.id, not Box.box_id)
    pull_events = PullEvent.query.filter_by(box_id=box.id)\
                                .order_by(PullEvent.timestamp.desc()).all()
    
    return render_template('box_logs.html', 
                         box=box, 
                         hardware_type=hardware_type,
                         lot_number=lot_number,
                         pull_events=pull_events)

@app.route('/export_excel')
def export_excel():
    """Export current inventory to Excel"""
    try:
        # Get filter parameters
        type_filter = request.args.get('type_filter', '')
        lot_filter = request.args.get('lot_filter', '')
        
        # Build query
        query = db.session.query(
            Box.box_id,
            HardwareType.name.label('type_name'),
            LotNumber.name.label('lot_name'),
            Box.box_number,
            Box.initial_quantity,
            Box.remaining_quantity,
            Box.barcode,
            Box.created_at
        ).join(HardwareType, Box.hardware_type_id == HardwareType.id)\
         .join(LotNumber, Box.lot_number_id == LotNumber.id)
        
        # Apply filters
        if type_filter:
            query = query.filter(HardwareType.name == type_filter)
        if lot_filter:
            query = query.filter(LotNumber.name == lot_filter)
        
        # Execute query
        results = query.order_by(Box.box_id).all()
        
        # Convert to DataFrame
        data = []
        for result in results:
            data.append({
                'Box ID': result.box_id,
                'Hardware Type': result.type_name,
                'Lot Number': result.lot_name,
                'Box Number': result.box_number,
                'Initial Quantity': result.initial_quantity,
                'Remaining Quantity': result.remaining_quantity,
                'Barcode': result.barcode,
                'Created Date': result.created_at.strftime('%Y-%m-%d %H:%M:%S') if result.created_at else ''
            })
        
        if not data:
            flash("No data to export", 'warning')
            return redirect(url_for('dashboard'))
        
        df = pd.DataFrame(data)
        
        # Create Excel file in memory
        output = io.BytesIO()
        with pd.ExcelWriter(output, engine='openpyxl') as writer:
            df.to_excel(writer, sheet_name='Inventory', index=False)
            
            # Get the workbook and worksheet
            workbook = writer.book
            worksheet = writer.sheets['Inventory']
            
            # Auto-adjust column widths
            for column in worksheet.columns:
                max_length = 0
                column_letter = column[0].column_letter
                for cell in column:
                    try:
                        if len(str(cell.value)) > max_length:
                            max_length = len(str(cell.value))
                    except:
                        pass
                adjusted_width = min(max_length + 2, 50)
                worksheet.column_dimensions[column_letter].width = adjusted_width
        
        output.seek(0)
        
        # Create response
        response = make_response(output.getvalue())
        response.headers['Content-Type'] = 'application/vnd.openxmlformats-officedocument.spreadsheetml.sheet'
        response.headers['Content-Disposition'] = f'attachment; filename=inventory_export_{datetime.now().strftime("%Y%m%d_%H%M%S")}.xlsx'
        
        return response
        
    except Exception as e:
        app.logger.error(f"Error exporting Excel: {str(e)}")
        flash("An error occurred while exporting to Excel", 'error')
        return redirect(url_for('dashboard'))

@app.route('/get_box_info/<barcode>')
def get_box_info(barcode):
    """API endpoint to get box info by barcode"""
    box = Box.query.filter_by(barcode=barcode).first()
    if box:
        hardware_type = HardwareType.query.get(box.hardware_type_id)
        lot_number = LotNumber.query.get(box.lot_number_id)
        return jsonify({
            'found': True,
            'box_id': box.box_id,
            'hardware_type': hardware_type.name,
            'lot_number': lot_number.name,
            'remaining_quantity': box.remaining_quantity
        })
    else:
        return jsonify({'found': False})

@app.route('/manage_boxes')
@admin_required
def manage_boxes():
    """List and manage all boxes with grouping"""
    # Get filter parameters
    type_filter = request.args.get('type_filter', '')
    lot_filter = request.args.get('lot_filter', '')
    search_query = request.args.get('search', '')
    
    # Use DRY filter helper
    query = build_filtered_query(type_filter, lot_filter, search_query)
    
    # Order by type name first, then lot name, then box number
    results = query.order_by(HardwareType.name, LotNumber.name, Box.box_number).all()
    
    # Group boxes using helper function
    grouped_data = group_boxes_by_type_lot(results)
    
    # Calculate statistics with pre-computed lot counts for template
    total_stats = calculate_inventory_stats(grouped_data)
    
    # Pre-compute lot counts for each type code to avoid Jinja complexity
    type_code_stats = {}
    for type_code, type_lot_groups in grouped_data.items():
        type_code_stats[type_code] = {
            'lot_count': len(type_lot_groups),
            'total_boxes': sum(g.get('box_count', 0) for g in type_lot_groups.values())
        }
    
    # Get filter options
    types = HardwareType.query.order_by(HardwareType.name).all()
    lots = LotNumber.query.order_by(LotNumber.name).all()
    
    return render_template('manage_boxes.html',
                         grouped_data=grouped_data,
                         total_stats=total_stats,
                         type_code_stats=type_code_stats,
                         types=types,
                         lots=lots,
                         type_filter=type_filter,
                         lot_filter=lot_filter,
                         search_query=search_query)

@app.route('/edit_box/<int:box_id>', methods=['GET', 'POST'])
@admin_required
def edit_box(box_id):
    """Edit an existing box - FULL ADMIN ACCESS"""
    box = Box.query.get_or_404(box_id)
    hardware_type = HardwareType.query.get(box.hardware_type_id)
    lot_number = LotNumber.query.get(box.lot_number_id)
    
    if request.method == 'POST':
        try:
            # Get form data - ADMIN CAN EDIT EVERYTHING
            hardware_type_name = request.form.get('hardware_type', '').strip()
            new_hardware_type = request.form.get('new_hardware_type', '').strip()
            lot_number_name = request.form.get('lot_number', '').strip()
            new_lot_number = request.form.get('new_lot_number', '').strip()
            new_box_number = request.form.get('box_number', '').strip()
            new_barcode = request.form.get('barcode', '').strip()
            new_initial_quantity = request.form.get('initial_quantity', 0)
            new_current_quantity = request.form.get('current_quantity', 0)
            
            # Validation
            errors = []
            
            # Handle hardware type
            target_hardware_type = None
            if new_hardware_type:
                existing_type = HardwareType.query.filter_by(name=new_hardware_type).first()
                if existing_type:
                    target_hardware_type = existing_type
                else:
                    target_hardware_type = HardwareType()
                    target_hardware_type.name = new_hardware_type
                    db.session.add(target_hardware_type)
                    db.session.flush()
            elif hardware_type_name:
                target_hardware_type = HardwareType.query.filter_by(name=hardware_type_name).first()
                if not target_hardware_type:
                    errors.append("Invalid hardware type selected")
            else:
                errors.append("Hardware type is required")
            
            # Handle lot number  
            target_lot_number = None
            if new_lot_number:
                existing_lot = LotNumber.query.filter_by(name=new_lot_number).first()
                if existing_lot:
                    target_lot_number = existing_lot
                else:
                    target_lot_number = LotNumber()
                    target_lot_number.name = new_lot_number
                    db.session.add(target_lot_number)
                    db.session.flush()
            elif lot_number_name:
                target_lot_number = LotNumber.query.filter_by(name=lot_number_name).first()
                if not target_lot_number:
                    errors.append("Invalid lot number selected")
            else:
                errors.append("Lot number is required")
            
            # Validate other fields
            if not new_box_number:
                errors.append("Box number is required")
            
            if not new_barcode:
                errors.append("Barcode is required")
            elif new_barcode != box.barcode:
                existing_box = Box.query.filter_by(barcode=new_barcode).first()
                if existing_box:
                    errors.append("Barcode already exists")
            
            try:
                new_initial_quantity = int(new_initial_quantity)
                new_current_quantity = int(new_current_quantity)
                if new_initial_quantity <= 0:
                    errors.append("Initial quantity must be greater than 0")
                if new_current_quantity < 0:
                    errors.append("Current quantity cannot be negative")
                if new_current_quantity > new_initial_quantity:
                    errors.append("Current quantity cannot exceed initial quantity")
            except (ValueError, TypeError):
                errors.append("Quantities must be valid numbers")
            
            if errors:
                for error in errors:
                    flash(error, 'error')
                form_data = request.form.to_dict()
                types = HardwareType.query.all()
                lots = LotNumber.query.all()
                return render_template('edit_box.html', box=box, hardware_type=hardware_type, 
                                     lot_number=lot_number, form_data=form_data, types=types, lots=lots)
            
            # Update box with new values
            if target_hardware_type:
                box.hardware_type_id = target_hardware_type.id
            if target_lot_number:
                box.lot_number_id = target_lot_number.id
            box.box_number = new_box_number
            box.barcode = new_barcode
            box.initial_quantity = new_initial_quantity
            box.remaining_quantity = new_current_quantity
            
            # Regenerate box_id
            if target_hardware_type and target_lot_number:
                box.box_id = generate_box_id(target_hardware_type.name, target_lot_number.name, new_box_number)
            
            db.session.commit()
            
            # Log the box edit action
            admin_user = session.get('admin_username', 'Unknown Admin')
            log_action(
                action_type='box_edit',
                user=admin_user,
                box_id=box.box_id,
                hardware_type=target_hardware_type.name if target_hardware_type else None,
                lot_number=target_lot_number.name if target_lot_number else None,
                details={
                    'new_initial_quantity': new_initial_quantity,
                    'new_current_quantity': new_current_quantity,
                    'new_barcode': new_barcode
                }
            )
            
            flash(f"Box {box.box_id} updated successfully!", 'success')
            return redirect(url_for('manage_boxes'))
            
        except Exception as e:
            db.session.rollback()
            app.logger.error(f"Error updating box: {str(e)}")
            flash("An error occurred while updating the box", 'error')
    
    # GET request - show form
    types = HardwareType.query.all()
    lots = LotNumber.query.all()
    return render_template('edit_box.html', box=box, hardware_type=hardware_type, 
                         lot_number=lot_number, types=types, lots=lots)

@app.route('/delete_box/<int:box_id>', methods=['POST'])
@admin_required
def delete_box(box_id):
    """Delete a box and all its pull events"""
    try:
        box = Box.query.get_or_404(box_id)
        box_id_name = box.box_id
        
        # Check if box has any pull events
        pull_events_count = PullEvent.query.filter_by(box_id=box_id).count()
        
        # Delete all pull events first (due to foreign key constraint)
        PullEvent.query.filter_by(box_id=box_id).delete()
        
        # Get box details for logging before deletion
        hardware_type = HardwareType.query.get(box.hardware_type_id)
        lot_number = LotNumber.query.get(box.lot_number_id)
        
        # Delete the box
        db.session.delete(box)
        db.session.commit()
        
        # Log the box deletion action
        admin_user = session.get('admin_username', 'Unknown Admin')
        log_action(
            action_type='box_delete',
            user=admin_user,
            box_id=box_id_name,
            hardware_type=hardware_type.name if hardware_type else None,
            lot_number=lot_number.name if lot_number else None,
            details={
                'pull_events_deleted': pull_events_count,
                'initial_quantity': box.initial_quantity,
                'remaining_quantity': box.remaining_quantity
            }
        )
        
        flash(f"Box {box_id_name} and {pull_events_count} pull events deleted successfully!", 'success')
        
    except Exception as e:
        db.session.rollback()
        app.logger.error(f"Error deleting box: {str(e)}")
        flash("An error occurred while deleting the box", 'error')
    
    return redirect(url_for('manage_boxes'))

@app.route('/admin/action_log')
@admin_required
def action_log():
    """Admin-only action log page"""
    from collections import Counter
    
    # Get filter parameters
    action_type_filter = request.args.get('action_type', '')
    user_filter = request.args.get('user', '')
    
    # Build query
    query = ActionLog.query
    
    # Apply filters
    if action_type_filter:
        query = query.filter(ActionLog.action_type == action_type_filter)
    if user_filter:
        query = query.filter(ActionLog.user.ilike(f'%{user_filter}%'))
    
    # Order by most recent first
    action_logs = query.order_by(ActionLog.timestamp.desc()).limit(500).all()
    
    # Calculate summary counts
    counts = Counter(log.action_type for log in action_logs)
    
    # Get unique action types and users for filter dropdowns
    action_types = db.session.query(ActionLog.action_type).distinct().all()
    action_types = [at[0] for at in action_types]
    
    users = db.session.query(ActionLog.user).distinct().all()
    users = [u[0] for u in users]
    
    return render_template('admin_action_log.html', 
                         action_logs=action_logs,
                         action_types=action_types,
                         users=users,
                         counts=counts,
                         action_type_filter=action_type_filter,
                         user_filter=user_filter)

@app.route('/export_all_logs')
@admin_required
def export_all_logs():
    """Export all inventory event logs with filters"""
    # Get filter parameters
    date_from = request.args.get('date_from', '')
    date_to = request.args.get('date_to', '')
    box_type_filter = request.args.get('box_type', '')
    lot_filter = request.args.get('lot', '')
    format_type = request.args.get('format', 'excel')  # excel or csv
    
    # Build query for pull events
    query = db.session.query(
        PullEvent, Box, HardwareType, LotNumber
    ).join(Box, PullEvent.box_id == Box.id
    ).join(HardwareType, Box.hardware_type_id == HardwareType.id
    ).join(LotNumber, Box.lot_number_id == LotNumber.id)
    
    # Apply filters
    if date_from:
        try:
            from_date = datetime.strptime(date_from, '%Y-%m-%d')
            query = query.filter(PullEvent.timestamp >= from_date)
        except ValueError:
            pass
    
    if date_to:
        try:
            to_date = datetime.strptime(date_to, '%Y-%m-%d')
            # Add one day to include the entire end date
            to_date = to_date.replace(hour=23, minute=59, second=59)
            query = query.filter(PullEvent.timestamp <= to_date)
        except ValueError:
            pass
    
    if box_type_filter:
        query = query.filter(HardwareType.name.ilike(f'%{box_type_filter}%'))
    
    if lot_filter:
        query = query.filter(LotNumber.name.ilike(f'%{lot_filter}%'))
    
    # Order by timestamp
    results = query.order_by(PullEvent.timestamp.desc()).all()
    
    # Prepare data for export
    export_data = []
    for pull_event, box, hardware_type, lot_number in results:
        # Calculate previous quantity (current + quantity change)
        previous_qty = box.remaining_quantity + abs(pull_event.quantity)
        event_type = "Return" if pull_event.quantity > 0 else "Pull"
        
        export_data.append({
            'Box ID': box.box_id,
            'Type': hardware_type.name,
            'Lot': lot_number.name,
            'MO': pull_event.mo or '',
            'Previous Quantity': previous_qty,
            'Event Type': event_type,
            'Quantity': abs(pull_event.quantity),
            'Operator': pull_event.operator or '',
            'QC Personnel': pull_event.qc_personnel or '',
            'Timestamp': pull_event.timestamp.strftime('%Y-%m-%d %H:%M:%S')
        })
    
    if format_type == 'csv':
        # Generate CSV
        output = io.StringIO()
        if export_data:
            writer = csv.DictWriter(output, fieldnames=export_data[0].keys())
            writer.writeheader()
            writer.writerows(export_data)
        
        response = make_response(output.getvalue())
        response.headers['Content-Type'] = 'text/csv'
        response.headers['Content-Disposition'] = f'attachment; filename=all_inventory_logs_{datetime.now().strftime("%Y%m%d_%H%M%S")}.csv'
        return response
    else:
        # Generate Excel
        df = pd.DataFrame(export_data)
        output = io.BytesIO()
        
        with pd.ExcelWriter(output, engine='openpyxl') as writer:
            df.to_excel(writer, sheet_name='Inventory Logs', index=False)
        
        output.seek(0)
        response = make_response(output.getvalue())
        response.headers['Content-Type'] = 'application/vnd.openxmlformats-officedocument.spreadsheetml.sheet'
        response.headers['Content-Disposition'] = f'attachment; filename=all_inventory_logs_{datetime.now().strftime("%Y%m%d_%H%M%S")}.xlsx'
        return response

@app.route('/export_box_history/<box_id>')
@admin_required
def export_box_history(box_id):
    """Export event history for a specific box"""
    box = Box.query.filter_by(box_id=box_id).first()
    if not box:
        flash("Box not found", 'error')
        return redirect(url_for('manage_boxes'))
    
    # Get pull events for this box
    pull_events = PullEvent.query.filter_by(box_id=box.id).order_by(PullEvent.timestamp.desc()).all()
    
    # Prepare data for export
    export_data = []
    current_qty = box.remaining_quantity
    
    for event in reversed(pull_events):  # Process in chronological order
        previous_qty = current_qty - event.quantity
        event_type = "Return" if event.quantity > 0 else "Pull"
        
        export_data.append({
            'Box ID': box.box_id,
            'Type': box.hardware_type.name,
            'Lot': box.lot_number.name,
            'MO': event.mo or '',
            'Previous Quantity': previous_qty,
            'Event Type': event_type,
            'Quantity': abs(event.quantity),
            'Operator': event.operator or '',
            'QC Personnel': event.qc_personnel or '',
            'Timestamp': event.timestamp.strftime('%Y-%m-%d %H:%M:%S')
        })
        current_qty = previous_qty
    
    # Reverse to show most recent first
    export_data.reverse()
    
    # Generate Excel
    df = pd.DataFrame(export_data)
    output = io.BytesIO()
    
    with pd.ExcelWriter(output, engine='openpyxl') as writer:
        df.to_excel(writer, sheet_name=f'Box_{box_id}_History', index=False)
    
    output.seek(0)
    response = make_response(output.getvalue())
    response.headers['Content-Type'] = 'application/vnd.openxmlformats-officedocument.spreadsheetml.sheet'
    response.headers['Content-Disposition'] = f'attachment; filename=box_{box_id}_history_{datetime.now().strftime("%Y%m%d_%H%M%S")}.xlsx'
    return response

@app.route('/generate_barcode_api')
def generate_barcode_api():
    """API endpoint to generate a new barcode"""
    # Check if custom barcode is provided
    custom_barcode = request.args.get('barcode', '').strip()
    barcode_type = request.args.get('type', 'qrcode')  # Default to QR code
    
    if custom_barcode:
        # Use provided barcode
        barcode_data = custom_barcode
    else:
        # Generate new unique barcode
        barcode_data = generate_unique_barcode()
    
    barcode_image = generate_barcode_image(barcode_data, barcode_type)
    
    return jsonify({
        'barcode': barcode_data,
        'image': barcode_image,
        'type': barcode_type
    })

@app.route('/print_template')
def print_template():
    """Barcode template page for printing"""
    # Get barcode data from query parameters
    barcodes = request.args.getlist('barcodes')
    box_ids = request.args.getlist('box_ids')
    barcode_type = request.args.get('type', 'qrcode')  # Default to QR code
    
    # If specific box IDs provided, get their barcodes
    if box_ids:
        boxes = Box.query.filter(Box.box_id.in_(box_ids)).all()
        barcode_data = []
        for box in boxes:
            barcode_data.append({
                'barcode': box.barcode,
                'box_id': box.box_id,
                'hardware_type': box.hardware_type.name if box.hardware_type else 'N/A',
                'lot_number': box.lot_number.name if box.lot_number else 'N/A',
                'remaining_quantity': box.remaining_quantity,
                'image': generate_barcode_image(box.barcode, barcode_type)
            })
    elif barcodes:
        # Custom barcodes provided
        barcode_data = []
        for barcode in barcodes:
            barcode_data.append({
                'barcode': barcode,
                'box_id': 'N/A',
                'hardware_type': 'Custom Barcode',
                'lot_number': 'N/A',
                'remaining_quantity': 0,
                'image': generate_barcode_image(barcode, barcode_type)
            })
    else:
        barcode_data = []
    
    return render_template('print_template.html', barcode_data=barcode_data)

@app.route('/bulk_print_barcodes', methods=['GET', 'POST'])
@admin_required
def bulk_print_barcodes():
    """Handle bulk barcode printing with GET fallback"""
    
    if request.method == 'GET':
        # Graceful fallback for accidental GET requests
        flash("Please select boxes from the dashboard to print barcodes", 'info')
        return redirect(url_for('dashboard'))
    
    # Original POST logic
    selected_boxes = request.form.getlist('selected_boxes')
    barcode_type = request.form.get('barcode_type', 'qrcode')
    
    if not selected_boxes:
        flash("No boxes selected for printing", 'warning')
        return redirect(url_for('dashboard'))
    
    # Redirect to print template with selected box IDs and barcode type
    return redirect(url_for('print_template', box_ids=selected_boxes, type=barcode_type))

@app.route('/bulk_print_logs', methods=['GET', 'POST'])
@admin_required
def bulk_print_logs():
    """Handle bulk log printing with GET fallback"""
    
    if request.method == 'GET':
        # Graceful fallback for accidental GET requests
        flash("Please select boxes from the dashboard to print logs", 'info')
        return redirect(url_for('dashboard'))
    
    # Original POST logic follows
    selected_boxes = request.form.getlist('selected_boxes')
    if not selected_boxes:
        flash("No boxes selected for log printing", 'warning')
        return redirect(url_for('dashboard'))
    
    # Generate combined log export for selected boxes
    boxes = Box.query.filter(Box.box_id.in_(selected_boxes)).all()
    if not boxes:
        flash("Selected boxes not found", 'error')
        return redirect(url_for('dashboard'))
    
    # Get pull events for selected boxes
    box_ids = [box.id for box in boxes]
    pull_events = db.session.query(
        PullEvent, Box, HardwareType, LotNumber
    ).join(Box, PullEvent.box_id == Box.id
    ).join(HardwareType, Box.hardware_type_id == HardwareType.id
    ).join(LotNumber, Box.lot_number_id == LotNumber.id
    ).filter(Box.id.in_(box_ids)
    ).order_by(Box.box_id, PullEvent.timestamp.desc()).all()
    
    # Prepare data for export
    export_data = []
    for pull_event, box, hardware_type, lot_number in pull_events:
        previous_qty = box.remaining_quantity + abs(pull_event.quantity)
        event_type = "Return" if pull_event.quantity > 0 else "Pull"
        
        export_data.append({
            'Box ID': box.box_id,
            'Type': hardware_type.name,
            'Lot': lot_number.name,
            'MO': pull_event.mo or '',
            'Previous Quantity': previous_qty,
            'Event Type': event_type,
            'Quantity': abs(pull_event.quantity),
            'Operator': pull_event.operator or '',
            'QC Personnel': pull_event.qc_personnel or '',
            'Timestamp': pull_event.timestamp.strftime('%Y-%m-%d %H:%M:%S')
        })
    
    # Generate Excel
    df = pd.DataFrame(export_data)
    output = io.BytesIO()
    
    with pd.ExcelWriter(output, engine='openpyxl') as writer:
        df.to_excel(writer, sheet_name='Selected_Boxes_Logs', index=False)
    
    output.seek(0)
    response = make_response(output.getvalue())
    response.headers['Content-Type'] = 'application/vnd.openxmlformats-officedocument.spreadsheetml.sheet'
    response.headers['Content-Disposition'] = f'attachment; filename=selected_boxes_logs_{datetime.now().strftime("%Y%m%d_%H%M%S")}.xlsx'
    return response

@app.route('/print_box_barcode/<box_id>')
@admin_required
def print_box_barcode(box_id):
    """Print barcode for a specific box"""
    barcode_type = request.args.get('type', 'qrcode')
    return redirect(url_for('print_template', box_ids=[box_id], type=barcode_type))

@app.route('/print_box_history/<box_id>')
@admin_required
def print_box_history(box_id):
    """Print history for a specific box"""
    return redirect(url_for('export_box_history', box_id=box_id))

@app.route('/export_word', methods=['POST'])
def export_word():
    """Export barcode layout to Word document with real images and borders"""
    try:
        barcode_data = json.loads(request.form.get('barcode_data', '[]'))
        template_data = json.loads(request.form.get('template', '{}'))
        
        # Create Word document
        doc = Document()
        
        # Set document orientation and margins (remove redundant orientation line)
        section = doc.sections[0]
        section.page_width = Inches(8.5)  # Letter width
        section.page_height = Inches(11)  # Letter height
        section.left_margin = Inches(0.5)
        section.right_margin = Inches(0.5)
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)
        
        # No title - removed as requested
        
        # Get template settings
        columns = template_data.get('columns', 2)
        rows = template_data.get('rows', 5)
        
        # Add space for company logo at top
        logo_para = doc.add_paragraph()
        logo_para.add_run('\n' * 3)  # Space for logo
        
        # Create table with proper dimensions and spacing
        table = doc.add_table(rows=rows, cols=columns)
        table.style = None  # Remove default style to avoid shared borders
        
        # Set table spacing to separate labels
        table.autofit = False
        
        # Fill table with barcode data using improved approach
        for i, item_data in enumerate(barcode_data):
            if i >= rows * columns:
                break
                
            # Use cleaner cell access
            cell = table.rows[i // columns].cells[i % columns]
            
            # Clear cell and add proper spacing
            cell.paragraphs[0].clear()
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            
            # Add top padding
            p.add_run('\n')
            
            # 1) Try to embed the real barcode image with correct path handling
            dataset = item_data.get('dataset', {})
            barcode_code = item_data.get('barcode', '').replace('Barcode: ', '')
            
            # Try to get barcode from the actual data structure
            if not barcode_code and 'boxId' in dataset:
                # Use box ID as barcode identifier if no direct barcode
                barcode_code = dataset['boxId']
            
            img_found = False
            if barcode_code:
                # Try multiple image paths and formats
                possible_paths = [
                    f"static/barcodes/{barcode_code}_qrcode.png",
                    f"static/barcodes/{barcode_code}_code128.png", 
                    f"static/barcodes/{barcode_code}.png"
                ]
                
                for img_path in possible_paths:
                    if os.path.exists(img_path):
                        try:
                            run = p.add_run()
                            run.add_picture(img_path, width=Inches(1.8))
                            img_found = True
                            break
                        except Exception as img_err:
                            continue
            
            # Fallback if no image found
            if not img_found:
                img_src = item_data.get('image_url', '')
                try:
                    if img_src.startswith('data:'):
                        header, data = img_src.split(',', 1)
                        img_bytes = io.BytesIO(base64.b64decode(data))
                        run = p.add_run()
                        run.add_picture(img_bytes, width=Inches(1.8))
                        img_found = True
                    elif img_src.startswith('/static/'):
                        img_path = img_src.lstrip('/')
                        if os.path.exists(img_path):
                            run = p.add_run()
                            run.add_picture(img_path, width=Inches(1.8))
                            img_found = True
                except Exception:
                    pass
            
            # If still no image, add placeholder space
            if not img_found:
                p.add_run('[QR Code Space]')
            
            # 2) Add spacing and data fields
            p.add_run('\n\n')  # Space between image and data
            
            # Add dataset information (already defined above)
            text_fields = item_data.get('text_fields', {})
            
            # Add barcode code
            if barcode_code:
                p.add_run(f"Code: {barcode_code}\n")
            
            # Add dataset information with proper field mapping
            for label, key in (('Type', 'type'), ('Lot', 'lot'), 
                              ('Qty', 'quantity'), ('Box ID', 'boxId')):
                val = dataset.get(key)
                if val:
                    p.add_run(f"{label}: {val}\n")
            
            # Add custom text fields  
            for field_type, value in text_fields.items():
                if value and value.strip():
                    p.add_run(f"{field_type.title()}: {value}\n")
            
            # Add bottom padding
            p.add_run('\n')
            
            # 3) Apply individual borders (no shared borders)
            set_cell_border(cell)
        
        # Save to BytesIO
        doc_buffer = io.BytesIO()
        doc.save(doc_buffer)
        doc_buffer.seek(0)
        
        # Create filename
        timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
        filename = f'barcode_layout_{timestamp}.docx'
        
        return make_response(
            doc_buffer.getvalue(),
            200,
            {
                'Content-Type': 'application/vnd.openxmlformats-officedocument.wordprocessingml.document',
                'Content-Disposition': f'attachment; filename="{filename}"'
            }
        )
        
    except Exception as e:
        flash(f'Error generating Word document: {str(e)}', 'error')
        return redirect(url_for('print_template'))

# Create tables
with app.app_context():
    db.create_all()
    
@app.route('/healthz')
def health_check():
    return 'OK', 200

if __name__ == '__main__':
    app.run(host='0.0.0.0', port=5000, debug=True)
